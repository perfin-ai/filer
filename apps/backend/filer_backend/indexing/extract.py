"""Text extraction by file type. Unsupported/binary types yield ("", "skipped")."""

import csv
from pathlib import Path

MAX_CHARS = 1_000_000  # cap extracted text so a huge file can't blow up memory

_TEXT_EXTS = {
    "txt", "md", "markdown", "rst", "log", "json", "yaml", "yml", "ini", "toml",
    "py", "js", "ts", "html", "xml", "tex",
}


def _pdf(path: Path) -> str:
    import pypdfium2 as pdfium

    parts: list[str] = []
    pdf = pdfium.PdfDocument(str(path))
    try:
        for page in pdf:
            tp = page.get_textpage()
            try:
                parts.append(tp.get_text_range())
            finally:
                tp.close()
            page.close()
            if sum(len(p) for p in parts) >= MAX_CHARS:
                break
    finally:
        pdf.close()
    return "\n".join(parts)


def _docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def _xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    rows: list[str] = []
    try:
        for ws in wb.worksheets:
            rows.append(f"# {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append("\t".join(cells))
                if sum(len(r) for r in rows) >= MAX_CHARS:
                    break
    finally:
        wb.close()
    return "\n".join(rows)


def _delimited(path: Path) -> str:
    delim = "\t" if path.suffix.lower() == ".tsv" else ","
    out: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        for row in csv.reader(f, delimiter=delim):
            out.append("\t".join(row))
            if sum(len(r) for r in out) >= MAX_CHARS:
                break
    return "\n".join(out)


def _native(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")[:MAX_CHARS]


def extract_text(path: Path) -> tuple[str, str]:
    """Return (text, parser_used). ("", "skipped") for unsupported/binary types.

    Raises on a parser failure so the caller can mark extraction_status="failed".
    """
    ext = path.suffix.lower().lstrip(".")
    if ext == "pdf":
        return _pdf(path)[:MAX_CHARS], "pypdfium2"
    if ext == "docx":
        return _docx(path)[:MAX_CHARS], "python-docx"
    if ext in ("xlsx", "xlsm"):
        return _xlsx(path)[:MAX_CHARS], "openpyxl"
    if ext in ("csv", "tsv"):
        return _delimited(path)[:MAX_CHARS], "csv"
    if ext in _TEXT_EXTS or ext == "":
        return _native(path), "native"
    return "", "skipped"
